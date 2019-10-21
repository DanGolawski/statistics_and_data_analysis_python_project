from docx import Document
from docx.shared import Inches
from lxml import html
import requests
import math
import random
import matplotlib.pyplot as plt;

plt.rcdefaults()
import numpy as np
import matplotlib.pyplot as plt


def countTechnology(list):
    counter = 0
    for i in list:
        if i == 'Technology':
            counter += 1

    return counter


def decision(x):
    if x > -2.38:
        return "Nie odrzucam H0 - nie popełniłem błędu I rodzaju"
    else:
        return "Odrzucam H0 - popełniłem błąd I rodzaju (odrzucam prawdziwą hipotezę)"


def createFirstPlot(tabx, tabSec, title):
    x_tab = []
    y_tab = []
    for i in range(0, len(tabx)):
        if tabSec[i] == 'Technology':
            if tabx[i] not in x_tab:
                x_tab.append(tabx[i])
                y_tab.append(1)
            else:
                y_tab[x_tab.index(tabx[i])] += 1

    # bubble sort
    n = len(x_tab)
    for i in range(n):
        for j in range(0, n - i - 1):
            if y_tab[j] < y_tab[j + 1]:
                y_tab[j], y_tab[j + 1] = y_tab[j + 1], y_tab[j]
                x_tab[j], x_tab[j + 1] = x_tab[j + 1], x_tab[j]

    y_pos = np.arange(len(x_tab))
    plt.bar(y_pos, y_tab, align='center', alpha=0.5)
    plt.xticks(y_pos, x_tab)
    plt.xticks(rotation=90)
    plt.ylabel('LICZBA MILIARDERÓW')
    plt.title('ROZKŁAD LISTY 500 NAJBOGATSZYCH Z TECHNOLOGY NA KRAJE')
    plt.subplots_adjust(bottom=0.3)
    plt.savefig(title + '.png', dpi=200)
    return ("Z wykresu dostrzegamy, że największy odsetek miliarderów z listy TOP 500 dla branży IT (Technology) pochodzi z " +
            x_tab[0] + " (" + str(max(y_tab)) + '), a najmniejszy z ' +
            x_tab[-1] + " (" + str(min(y_tab)) + ").")


def createSecondPlot(tab, title, base):
    secTab = []
    secTabCount = []
    for i in range(0, len(tab)):
        if tab[i] not in secTab:
            secTab.append(tab[i])
            secTabCount.append(1)
        else:
            secTabCount[secTab.index(tab[i])] += 1

    #bubble sort
    n = len(secTab)
    for i in range(n):
        for j in range(0, n - i - 1):
            if secTabCount[j] < secTabCount[j + 1]:
                secTabCount[j], secTabCount[j + 1] = secTabCount[j + 1], secTabCount[j]
                secTab[j], secTab[j + 1] = secTab[j + 1], secTab[j]

    y_pos = np.arange(len(secTab))
    plt.bar(y_pos, secTabCount, align='center', alpha=0.5)
    plt.xticks(y_pos, secTab)
    plt.xticks(rotation=90)
    plt.ylabel('LICZBA MILIARDERÓW PER BRANŻA')
    plt.title('ROZKŁAD LISTY 500 NAJBOGATSZYCH NA BRANŻE W '+base)
    plt.subplots_adjust(bottom=0.3)
    plt.savefig(title + '.png', dpi=200)

    return "Z wykresu widzimy, że branża Technology zajmuje " + str(secTab.index('Technology')+1) + \
           ". miejsce w zestawieniu liczny miliarderów per branża."





page = requests.get('https://ceoworld.biz/2018/05/30/rich-list/')
tree = html.fromstring(page.content)

# Lista miliardereów
peopleList = tree.xpath('//*[@id="tablepress-303"]/tbody/tr/td[2]/text()')

# Branże miliarderów
sectorList = tree.xpath('//*[@id="tablepress-303"]/tbody/tr/td[5]/text()')

# lista krajów
countryList = tree.xpath('//*[@id="tablepress-303"]/tbody/tr/td[4]/text()')


# Tworzę wykresy wraz z komentarzami dla populacji
plotComment1 = createFirstPlot(countryList, sectorList, "wykres1")
plotComment2 = createSecondPlot(sectorList, "wykres2", "POPULACJI")





# liczba miliarderów z branży Technology w populacji
technologyInPopulation = countTechnology(sectorList)

# liczebność populacji
n = len(sectorList)

# średnia 'Technology' w populacji
mi = technologyInPopulation / len(sectorList)

# odchylenie standardowe w populacji
sigma = math.sqrt(mi * (1 - mi))

# tworzę próbkę danych - 100 elementów z 500
sample = random.sample(sectorList, 100)

# tworzę wykres i koentarz dla próbki
plotComment3 = createSecondPlot(sample, "wykres3", "PRÓBIE")

# liczę średnią wystąpień 'Technology' w próbce
sampleSuccess = countTechnology(sample)
mean = sampleSuccess / 100
# print(mean)

# wartość mi
my_mi = 0.1

# Obliczenia na test Z
z = (mean - my_mi) / (sigma / math.sqrt(n))



document = Document()

document.add_heading('Sprawozdanie SAD', 0)

myParagraph = ("Według rankingu 100 najbogatszych polaków szacuje się, że miliarderzy z branży technologicznej stanowią "+str(my_mi)+" wszystkich z listy. W związku z tym chciałbym sprawdzić hipotezę, że " \
                   "odsetek miliarderów z branży technologicznej na świecie jest większy od "+str(my_mi)+". Moją populację stanowi 500 światowych miliarderów, natomiast populacja to stuelementowy podzbiór " \
                   "tej listy, w którym algorytm za każdym razem dobiera losowe elementy. Skoro moja populacja jest znana, wiem, że postawiona hipoteza jest prawidłowa, a do testów wykorzystuję test \"Z\". " \
                   "Dane (obserwacyjne) zostały pobrane ze strony https://ceoworld.biz/2018/05/30/rich-list/. Zawierają informacje o majątku, kraju pochodzenia oraz branży miliarderów. " \
                   "W obliczeniach wykorzystuję rozkład Bernouliego - jeśli element jest związany z Technology, jest sukces (1), w przeciwnym wypadku porażka (0).")


p = document.add_paragraph(myParagraph)


document.add_picture('wykres1.png', width=Inches(5.25))
document.add_paragraph(plotComment1, style='List Bullet')

document.add_picture('wykres2.png', width=Inches(5.25))
document.add_paragraph(plotComment2, style='List Bullet')

document.add_picture('wykres3.png', width=Inches(5.25))
document.add_paragraph(plotComment3, style='List Bullet')

document.add_heading('Dane o populacji i próbce', level=1)
document.add_paragraph('Liczebność populacji n : ' + str(n), style='List Bullet')
document.add_paragraph('Liczba sukcesów w populacji : ' + str(technologyInPopulation), style='List Bullet')
document.add_paragraph('Odsetek "Technology" w populacji (p): ' + str(round(mi, 3)), style='List Bullet')
document.add_paragraph('Odchylenie standardowe w populacji σ : ' + str(round(sigma, 3)), style='List Bullet')
document.add_paragraph('Liczebność próbki : ' + str(len(sample)), style='List Bullet')
document.add_paragraph('Liczba sukcesów w próbce : ' + str(sampleSuccess), style='List Bullet')
document.add_paragraph('Odsetek "Technology" w próbce x̅ : ' + str(round(mean, 3)), style='List Bullet')
document.add_paragraph('Odsetek miliarderów związanych z sektorem Technology jest niemniejszy niż '+str(my_mi), style='List Bullet')
document.add_paragraph('Odsetek miliarderów związanych z sektorem Technology jest mniejszy niż '+str(my_mi), style='List Bullet')
document.add_paragraph('H0 : μ ≥ ' + str(my_mi) + (' (na podstwie wiedzy o populacji, wiem, że hipoteza ta jest prawdziwa)'), style='List Bullet')
document.add_paragraph('H0 : μ < ' + str(my_mi), style='List Bullet')
document.add_paragraph('Poziom ufności α : 0.05', style='List Bullet')
document.add_paragraph('Przedział ufności α=0.05 : (-∞, -2,38)', style='List Bullet')
document.add_paragraph('Wynik testu Z : ' + str(round(z, 3)), style='List Bullet')
document.add_paragraph('Decyzja o odrzuceniu H0 : ' + decision(z), style='List Bullet')


document.add_page_break()

document.save('sprawozdanie.docx')
